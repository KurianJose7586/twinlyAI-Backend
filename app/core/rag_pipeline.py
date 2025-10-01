import os
import json
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, AIMessage
from app.core.config import settings

def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1]

def json_to_text(json_data: dict) -> str:
    text = ""
    for key, value in json_data.items():
        if isinstance(value, dict):
            text += f"{key.replace('_', ' ').title()}:\n"
            for sub_key, sub_value in value.items():
                text += f"  {sub_key.replace('_', ' ').title()}: {sub_value}\n"
        elif isinstance(value, list):
            text += f"{key.replace('_', ' ').title()}:\n"
            for item in value:
                if isinstance(item, dict):
                    for item_key, item_value in item.items():
                        text += f"  - {item_key.replace('_', ' ').title()}: {item_value}\n"
                else:
                    text += f"- {item}\n"
        else:
            text += f"{key.replace('_', ' ').title()}: {value}\n"
    return text

def extract_text_from_file(file_path: Path) -> str:
    if file_path.suffix == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            return "".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_path.suffix == ".docx":
        doc = DocxDocument(file_path)
        return "\n".join(para.text for para in doc.paragraphs)
    elif file_path.suffix == ".txt":
        return file_path.read_text(encoding="utf-8")
    elif file_path.suffix == ".json":
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json_to_text(data)
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

class RAGPipeline:
    def __init__(self, bot_id: str, user_id: str, bot_name: str):
        self.bot_id = bot_id
        self.user_id = user_id
        self.bot_name = bot_name
        self.data_path = Path("data") / user_id / bot_id
        self.index_path = self.data_path / "faiss_index"
        
        self.embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5")
        
        self.llm = ChatGroq(
            model_name="llama3-8b-8192", 
            temperature=0.7, 
            groq_api_key=settings.GROQ_API_KEY
        )
        
        self.vector_store = self._load_vector_store()
        self.retrieval_chain = self._create_retrieval_chain()

    def _load_vector_store(self):
        if self.index_path.exists():
            try:
                return FAISS.load_local(
                    str(self.index_path), 
                    self.embeddings, 
                    allow_dangerous_deserialization=True
                )
            except Exception as e:
                print(f"Error loading vector store: {e}")
                return None
        return None

    def _create_retrieval_chain(self):
        if not self.vector_store:
            return None
        
        system_prompt = f"""
You are "{self.bot_name}," a professional AI assistant. Your task is to answer questions about a person based on their resume provided in the context.
**Persona & Introduction:**
- Your name is "{self.bot_name}".
- **Introduce yourself ONLY IF it is the first turn of the conversation or if asked "who are you?".**
- Your introduction should be: "Hello, I am {self.bot_name}, an AI assistant for [Person's Name]. I can answer questions based on their resume. How can I help?"
- You must extract the [Person's Name] from the context.
- Always speak about the person in the third person (e.g., "He has experience in...").
**Response Guidelines:**
- Answer exclusively from the <context>.
- If the information isn't in the context, politely state that.
- Use Markdown (bolding, bullet points) for clarity.
- For personal or off-topic questions, state that you can only answer professional questions based on the resume.
- DO NOT INCLUDE THE <think> </think> part in your resposnses.

<context>
{{context}}
</context>
"""
        prompt = ChatPromptTemplate.from_messages(
            [
                ("system", system_prompt),
                MessagesPlaceholder(variable_name="chat_history"),
                ("human", "{input}"),
            ]
        )
        
        question_answer_chain = create_stuff_documents_chain(self.llm, prompt)
        return create_retrieval_chain(self.vector_store.as_retriever(), question_answer_chain)

    async def load_and_index_document(self, file_path: str):
        text_content = extract_text_from_file(Path(file_path))
        documents = [Document(page_content=text_content)]
        
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(documents)

        self.data_path.mkdir(parents=True, exist_ok=True)
        self.vector_store = FAISS.from_documents(documents=splits, embedding=self.embeddings)
        self.vector_store.save_local(str(self.index_path))
        
        self.retrieval_chain = self._create_retrieval_chain()
        return True

    async def get_response_stream(self, user_message: str, chat_history: list = []):
        if not self.retrieval_chain:
            yield {"answer": "Error: The AI bot has not been properly initialized. Please upload a resume."}
            return
        
        # --- THIS IS THE FIX ---
        # Yield the entire chunk dictionary, not just the "answer" string.
        async for chunk in self.retrieval_chain.astream({
            "input": user_message,
            "chat_history": chat_history
        }):
            yield chunk
        # --- END OF FIX ---